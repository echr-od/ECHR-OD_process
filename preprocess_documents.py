import argparse
import json
import os
from os import listdir
from os.path import isfile, join
import re
import shutil
from docx import Document
from docx.shared import Inches
from docx.text.run import Run
import zipfile

TMP = '/tmp/echr_tmp_doc'

# Possible tags for each type of section
tags = {
    "SECTION_TITLE_STYLE": ['ECHR_Title_1', 'Ju_H_Head'],
    "HEADING_1_STYLE": ['ECHR_Heading_1', 'Ju_H_I_Roman'],
    "HEADING_2_STYLE" :['ECHR_Heading_2', 'Ju_H_A', 'Ju_H_a'],
    "HEADING_3_STYLE" :['ECHR_Heading_3', 'Ju_H_1.', 'Ju_H_1'],
    "HEADING_PARA": ['ECHR_Para', 'ECHR_Para_Quote', 'Ju_List',\
        'Ju_List_a', 'Ju_Para', 'Normal', 'Ju_Quot', 'Ju_H_Article',\
        'Ju_Para Char Char', 'Ju_Para Char', 'Ju_Para_Last', 'Opi_Para'],
    "DECISION_BODY": ['ECHR_Decision_Body', 'Ju_Judges']
}

# Different level of document to parse
levels = {
    "DECISION_BODY": -1,
    "SECTION_TITLE_STYLE": 1,
    "HEADING_1_STYLE": 2,
    "HEADING_2_STYLE": 3,
    "HEADING_3_STYLE": 4,
    "HEADING_PARA": 5
}

tag_to_level = {}
for k, v in tags.items():
    for t in v:
        tag_to_level[t] = levels[k]

OLD_PARSER_TAGS = ['header', 'Normal', 'Body Text 2', 'Body Text Indent 3', 'OldCommission', 'Heading 6', 'Heading 5', 'Heading 4']

internal_section_reference = {
    'procedure': "",
    'facts': "",
    'law': "",
    'conclusion': ""
}

def tag_elements(parsed):
    """Tag the elements in the parsed document.

        Tag the elements in the parsed documents
        according to some predifined sections.

        :param parsed: parsed document
        :type parsed: dict
        :return: parsed document with internal section references
        :rtype: dict
    """
    for i, section in parsed['elements']:
        pass # TODO: Tag the element in the parsed document
    parsed['elements'][-1]['section_name'] = 'conclusion'
    return parsed


def format_title(line):
    """Format title

        :param line: line to format as title
        :type line: str
        :return: formatted title
        :rtype: str
    """
    m = re.match(r'(\w+)\.(.+)', line)
    if m:
        return m.group(2).strip()
    else:
        return line

def parse_body(body):
    """Extract body members

        :param body: line to extract the body members from
        :type body: str
        :return: list of members with their role
        :rtype: [dict]
    """
    members = []
    body = body.replace('\nand ', '\n')
    body = body.replace('\t', '')
    body = body.split('\n')
    body = [b for b in body if len(b)]

    roles = []
    k = 0
    for i, t in enumerate(body):
        a = [j for j in t.split(',') if len(j)]
        members.append({'name': a[0]})
        if len(a) > 1:
            roles.append((k, i, a[1].lower().strip()))
            k = i + 1

    for r in roles:
        for i, m in enumerate(members[r[0]:r[1]+1]):
            members[r[0] + i]['role'] = r[2]

    return members


class Node:
    """Represent a rooted tree
    """

    def __init__(self, parent=None, level=0, content=None):
        self.parent = parent
        self.level = level
        self.content = content
        self.elements = []


def parse_document(doc):
    """Parse a document object to a tree

        :param doc: document object
        :type doc: Document
        :return: tree
        :rtype: Node
    """
    parsed = {}
    result = []
    section = {}

    decision_body = ""
    appender = Node() # Top level node
    for p in doc.paragraphs:
        line = p.text.strip()
        if not len(line):
            continue
        #print(p.style.name, p.text)
        level = tag_to_level.get(p.style.name, 0)
        if level > 0:
            if appender.level == 0 and not len(appender.elements) and level > 1:
                pass
                #print('HEADER')
            else:
                #print('L {} | App level: {}'.format(level, appender.level))
                if level < appender.level:
                    while(appender.level > level - 1):
                        appender = appender.parent
                elif level == appender.level:
                    appender = appender.parent
                node = Node(parent=appender, level=level, content=p.text)
                appender.elements.append(node)
                appender = node

        if level < 0:
            if level == -1:
                #print(p.text)
                decision_body += p.text
        #else:
        #    print(p.style.name, p.text)
    

    root = appender
    while(root.level != 0):
        root = root.parent

    def print_tree(root):
        """Utilitary function to print tree

            :param root: root of the tree
            :type root: Node
        """
        print("LEVEL {} {} {}".format(root.level, ' ' * root.level * 2, root.content.encode('utf-8') if root.content else 'ROOT'))
        if len(root.elements) == 0:
            return
        else:
            for e in root.elements:
                print_tree(e)

    def tree_to_json(root, res):
        """Recursively convert a tree into json

            :param root: root of the tree
            :type root: Node
            :param res: where to store result
            :type: res: dict
            :return: remaining tree
            :rtype: Node
        """
        node = {
            'content': root.content,
            'elements': []
        }
        for e in root.elements:
            node['elements'].append(tree_to_json(e, node))
        return node

    parsed = {'elements': []}
    parsed['elements'] = tree_to_json(root, parsed)['elements']
    #del parsed['elements']
    #for e in appender.elements:
    #    print(e.content)
    #print_tree(root)
    parsed['decision_body'] = parse_body(decision_body) if decision_body else []
    parsed = tag_elements(parsed)

    return parsed

PARSER = {
    'old': 'OLD',
    'new': 'NEW'
}

def format_paragraph(p):
    """Format paragraph

        :param line: line to format as title
        :type line: str
        :return: formatted title
        :rtype: str
    """
    match = re.search(r'^(?:\w+\.)(.+)', p)
    if match is not None:
        return match.group(1).strip()
    else:
        return p
    
def json_to_text_(doc, text_only=True, except_section=[]):
    res = []
    if not len(doc['elements']):
        res.append(format_paragraph(doc['content']))
    # text_only: remove the titles
    for e in doc['elements']:
        if not 'section_name' in e or e['section_name'] not in except_section:
            res.extend(json_to_text_(e, text_only=True, except_section=except_section))
    return res

def json_to_text(doc, text_only=True, except_section=[]):
    """Format json to text 

        :param doc: parsed document
        :type doc: dict
        :param text_only: return only text
        :type text_only: bool
        :param except_section: list of section to discard
        :type: except_section: list
        :return: textual representation of the document
        :rtype: str
    """
    return '\n'.join(json_to_text_(doc, text_only, except_section))

def select_parser(doc):
    """Select the parser to be used for a given document

        :param doc: document
        :type doc: Document
        :return: parser name
        :rtype: str
    """
    if all([True if p.style.name in OLD_PARSER_TAGS else False for p in doc.paragraphs]):
        return PARSER['old']
    else:
        return PARSER['new']

def main(args):
    input_file = os.path.join(args.build, 'cases_info/raw_cases_info_all.json')
    input_folder = os.path.join(args.build, 'raw_documents')
    output_folder = os.path.join(args.build, 'preprocessed_documents')
    if not args.u:
        try:
            if args.f:
                shutil.rmtree(output_folder)
            os.mkdir(output_folder)
        except Exception as e:
            print(e)
            exit(1)
            
    stats = {
        'parser_type':{
            'OLD': 0,
            'NEW': 0
        }
    }

    with open(input_file, 'r') as f:
        content = f.read()
        cases = json.loads(content)
        cases_index = {c['itemid']:i for i,c in enumerate(cases)}
        f.close()

    update = args.u
    correctly_parsed = 0
    failed = []
    files = [os.path.join(input_folder, f) for f in listdir(input_folder) if isfile(join(input_folder, f)) if '.docx' in f]
    print('# Process documents')
    for i, p in enumerate(files):
        id_doc = p.split('/')[-1].split('.')[0]
        print(' - Process document {} {}/{}'.format(id_doc, i, len(files)))
        filename_parsed = os.path.join(output_folder, '{}_parsed.json'.format(id_doc))
        if not update or not os.path.isfile(filename_parsed):
            try:
                p_ = update_docx(p)
                doc = Document(p_)
                parser = select_parser(doc)
                stats['parser_type'][parser] += 1
                if parser == 'NEW':
                    parsed = parse_document(doc)
                    parsed.update(cases[cases_index[id_doc]])
                    with open(os.path.join(output_folder, '{}_text_without_conclusion.txt'.format(id_doc)), 'w') as toutfile:
                        toutfile.write(json_to_text(parsed, True, ['conclusion']))
                    parsed['documents'] = ['{}.docx'.format(id_doc)]
                    parsed['content'] = {
                        '{}.docx'.format(id_doc): parsed['elements']
                    }
                    del parsed['elements']
                    with open(filename_parsed, 'w') as outfile:
                        json.dump(parsed, outfile, indent=4, sort_keys=True)
                    correctly_parsed += 1
                else:
                    raise Exception("OLD parser is not available yet.")
            except Exception as e:
                failed.append((id_doc,e))    
                print("\t->{} {}".format(p, e))
        else:
            print('\t-> Skip document because it is already processed')
            correctly_parsed += 1

    print('# Correctly parsed: {}/{} ({}%)'.format(correctly_parsed, len(files), (100. * correctly_parsed) / len(files)))
    print('# List of failed documents:')
    for e in failed:
        print('\t- {}: {}'.format(e[0], e[1]))


def update_docx(docname):
    """Update a docx such that it can be read by docx library.

        MSWord documents are a zip folder containing several XML files.
        As docx library cannot read 'smartTag', it is required to remove them.
        To do so, we open the zip, access the main XML file and manually sanitize it.

        :param docname: path to the document
        :type docname: str
        :return: path to the new document
        :rtype: str
    """
    # Remove temporary folder and files
    try:
        shutil.rmtree(TMP)
    except:
        pass

    try:
        os.remove('./_proxy.docx')
    except:
        pass

    # Extract the document
    zip_ref = zipfile.ZipFile(docname, 'r')
    zip_ref.extractall(TMP)
    zip_ref.close()

    # Sanitize
    with open(os.path.join(TMP, 'word/document.xml'), 'r') as file:
        content = file.read()
        lines = content.split('>')
        remove_open = True
        for i, l in enumerate(lines):
            if '<w:smartTag ' in l and remove_open:
                del lines[i]
                remove_open = False
            if '</w:smartTag'==l and not remove_open:
                del lines[i]
                remove_open = True
        file.close()
    content = '>'.join(lines)

    # Recompress the archive
    with open(os.path.join(TMP, 'word/document.xml'), 'w') as file:
        file.write(content)
    shutil.make_archive('./proxy', 'zip', TMP)

    output_file = './_proxy.docx'
    os.rename('./proxy.zip', output_file)

    try:
        os.remove('./_proxy.docx')
    except:
        pass
    
    return output_file


def parse_args(parser):
    args = parser.parse_args()

    # Check path
    return args

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Filter and format ECHR cases information')
    parser.add_argument('--build', type=str, default="./build/echr_database/")
    parser.add_argument('-f', action='store_true')
    parser.add_argument('-u', action='store_true')
    args = parse_args(parser)

    main(args)
 