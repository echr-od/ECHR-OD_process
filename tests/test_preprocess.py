from mock import patch
from docx import Document
import pytest
import json

from echr.steps.preprocess_documents import para_to_text, json_table_to_text, parse_document, load_judges_info


class TestPreprocessWord:

    @staticmethod
    @pytest.fixture
    def prepare():
        file = 'tests/data/judgments/001-83979.docx'
        expected_file = 'tests/data/judgments/001-83979_without_smarttags.docx'
        doc = Document(file)
        expected_doc = Document(expected_file)
        broken_list_paragraphs = []
        fixed_list_paragraphs = []
        expected_list_paragraphs = []

        for p in doc.paragraphs:
            broken_list_paragraphs.append(p.text)
            fixed_list_paragraphs.append(para_to_text(p))

        for p in expected_doc.paragraphs:
            expected_list_paragraphs.append(para_to_text(p))

        broken_list_paragraphs = [p for p in broken_list_paragraphs if p]
        fixed_list_paragraphs = [p for p in fixed_list_paragraphs if p]
        expected_list_paragraphs = [p for p in expected_list_paragraphs if p]
        return {'broken': broken_list_paragraphs, 'fixed': fixed_list_paragraphs, 'expected': expected_list_paragraphs}

    @staticmethod
    def test_len_of_paragraphs(prepare):
        assert len(prepare['fixed']) == len(prepare['expected'])

    @staticmethod
    def test_fixed_equals_expected(prepare):
        assert all([p == prepare['fixed'][i] for i, p in enumerate(prepare['expected'])])

    @staticmethod
    def test_broken_different_from_fixed(prepare):
        assert any([p != prepare['fixed'][i] for i, p in enumerate(prepare['broken'])])

    @staticmethod
    def test_broken_different_from_expected(prepare):
        assert any([p != prepare['expected'][i] for i, p in enumerate(prepare['broken'])])


class TestProcessTableAttachment:
    @staticmethod
    def test_json_table_to_text_emtpy():
        table = []
        expected = ""
        res = json_table_to_text(table)
        assert res == expected

    @staticmethod
    def test_json_table_to_text_text_only():
        table = [
            {'header1': 'val1', 'header2': 'val2'},
            {'header1': 'val3', 'header2': 'val4'}]
        expected = "header1 header2\nval1 val2\nval3 val4\n"
        res = json_table_to_text(table)
        assert res == expected

    @staticmethod
    def test_json_table_to_text_int_and_float():
        table = [
            {'header1': 'val1', 'header2': 10},
            {'header1': 12., 'header2': 'val4'}]
        expected = "header1 header2\nval1 10\n12.0 val4\n"
        res = json_table_to_text(table)
        assert res == expected


class TestParseDecisionBody:
    @staticmethod
    @pytest.fixture
    def prep():
        doc_ids = ['001-175180', '001-176769', '001-177299']
        files = ['tests/data/judgments/{doc_id}.docx'.format(doc_id=i) for i in doc_ids]
        docs = [Document(file) for file in files]
        build = './build/echr_database/'
        # inputs = [[docs[i], doc_ids[i], build] for i in range(len(doc_ids))]
        outputs = [
            [{'name': 'SAJÓ', 'info': {'end': '2017', 'full_name': 'András SAJÓ', 'start': '2008'}, 'role': 'judge'},
             {'name': 'KARAKAŞ', 'info': {'end': '2019', 'full_name': 'Işıl KARAKAŞ', 'start': '2008'},
              'role': 'judge'},
             {'name': 'NUSSBERGER', 'info': {'end': '2019', 'full_name': 'Angelika NUSSBERGER', 'start': '2011'},
              'role': 'judge'},
             {'name': 'HAJIYEV', 'info': {'end': '2016', 'full_name': 'Khanlar HAJIYEV', 'start': '2003'},
              'role': 'judge'},
             {'name': 'LÓPEZ GUERRA', 'info': {'end': '2018', 'full_name': 'Luis LÓPEZ GUERRA', 'start': '2008'},
              'role': 'judge'}, {'name': 'LAZAROVA TRAJKOVSKA',
                                 'info': {'end': '2017', 'full_name': 'Mirjana LAZAROVA TRAJKOVSKA', 'start': '2008'},
                                 'role': 'judge'},
             {'name': 'VUČINIĆ', 'info': {'end': '2018', 'full_name': 'Nebojša VUČINIĆ', 'start': '2008'},
              'role': 'judge'},
             {'name': 'DE GAETANO', 'info': {'end': '2019', 'full_name': 'Vincent A. DE GAETANO', 'start': '2010'},
              'role': 'judge'},
             {'name': 'POTOCKI', 'info': {'end': '2020', 'full_name': 'André POTOCKI', 'start': '2011'},
              'role': 'judge'},
             {'name': 'MAHONEY', 'info': {'end': '2016', 'full_name': 'Paul MAHONEY', 'start': '2012'},
              'role': 'judge'},
             {'name': 'VEHABOVIĆ', 'info': {'end': None, 'full_name': 'Faris VEHABOVIĆ', 'start': '2012'},
              'role': 'judge'},
             {'name': 'KŪRIS', 'info': {'end': '2004', 'full_name': 'Pranas KŪRIS', 'start': '1994'}, 'role': 'judge'},
             {'name': 'MOTOC', 'info': {'end': None, 'full_name': 'Iulia Antoanella MOTOC', 'start': '2014'},
              'role': 'judge'},
             {'name': 'KJØLBRO', 'info': {'end': None, 'full_name': 'Jon Fridrik KJØLBRO', 'start': '2014'},
              'role': 'judge'},
             {'name': 'MITS', 'info': {'end': None, 'full_name': 'Mārtiņš MITS', 'start': '2015'}, 'role': 'judge'},
             {'name': 'MOUROU VIKSTRÖM',
              'info': {'end': None, 'full_name': 'Stéphanie MOUROU-VIKSTRÖM', 'start': '2015'}, 'role': 'judge'},
             {'name': 'KUCSKO STADLMAYER',
              'info': {'end': None, 'full_name': 'Gabriele KUCSKO-STADLMAYER', 'start': '2015'}, 'role': 'judge'}],
            [{'name': 'RAIMONDI', 'info': {'end': '2019', 'full_name': 'Guido RAIMONDI', 'start': '2010'},
              'role': 'judge'},
             {'name': 'NUSSBERGER', 'info': {'end': '2019', 'full_name': 'Angelika NUSSBERGER', 'start': '2011'},
              'role': 'judge'}, {'name': 'LAZAROVA TRAJKOVSKA',
                                 'info': {'end': '2017', 'full_name': 'Mirjana LAZAROVA TRAJKOVSKA', 'start': '2008'},
                                 'role': 'judge'},
             {'name': 'LÓPEZ GUERRA', 'info': {'end': '2018', 'full_name': 'Luis LÓPEZ GUERRA', 'start': '2008'},
              'role': 'judge'},
             {'name': 'SAJÓ', 'info': {'end': '2017', 'full_name': 'András SAJÓ', 'start': '2008'}, 'role': 'judge'},
             {'name': 'KARAKAŞ', 'info': {'end': '2019', 'full_name': 'Işıl KARAKAŞ', 'start': '2008'},
              'role': 'judge'},
             {'name': 'PARDALOS', 'info': {'end': '2018', 'full_name': 'Kristina PARDALOS', 'start': '2009'},
              'role': 'judge'},
             {'name': 'POTOCKI', 'info': {'end': '2020', 'full_name': 'André POTOCKI', 'start': '2011'},
              'role': 'judge'},
             {'name': 'GRIŢCO', 'info': {'end': None, 'full_name': 'Valeriu GRIŢCO', 'start': '2012'}, 'role': 'judge'},
             {'name': 'VEHABOVIĆ', 'info': {'end': None, 'full_name': 'Faris VEHABOVIĆ', 'start': '2012'},
              'role': 'judge'},
             {'name': 'TURKOVIĆ', 'info': {'end': None, 'full_name': 'Ksenija TURKOVIĆ', 'start': '2013'},
              'role': 'judge'},
             {'name': 'LUBARDA', 'info': {'end': None, 'full_name': 'Branko LUBARDA', 'start': '2015'},
              'role': 'judge'},
             {'name': 'GROZEV', 'info': {'end': None, 'full_name': 'Yonko GROZEV', 'start': '2015'}, 'role': 'judge'},
             {'name': 'O’LEARY', 'info': {'end': None, 'full_name': 'Síofra O’LEARY', 'start': '2015'},
              'role': 'judge'},
             {'name': 'RANZONI', 'info': {'end': None, 'full_name': 'Carlo RANZONI (Swiss)', 'start': '2015'},
              'role': 'judge'}, {'name': 'MOUROU VIKSTRÖM',
                                 'info': {'end': None, 'full_name': 'Stéphanie MOUROU-VIKSTRÖM', 'start': '2015'},
                                 'role': 'judge'},
             {'name': 'KOSKELO', 'info': {'end': None, 'full_name': 'Pauliine KOSKELO', 'start': '2016'},
              'role': 'judge'}],
            [{'name': 'RAIMONDI', 'info': {'end': '2019', 'full_name': 'Guido RAIMONDI', 'start': '2010'},
              'role': 'judge'},
             {'name': 'SICILIANOS', 'info': {'end': '2021', 'full_name': 'Linos-Alexandre SICILIANOS', 'start': '2011'},
              'role': 'judge'},
             {'name': 'SPANO', 'info': {'end': None, 'full_name': 'Robert SPANO', 'start': '2013'}, 'role': 'judge'},
             {'name': 'LAZAROVA TRAJKOVSKA',
              'info': {'end': '2017', 'full_name': 'Mirjana LAZAROVA TRAJKOVSKA', 'start': '2008'}, 'role': 'judge'},
             {'name': 'HAJIYEV', 'info': {'end': '2016', 'full_name': 'Khanlar HAJIYEV', 'start': '2003'},
              'role': 'judge'},
             {'name': 'LÓPEZ GUERRA', 'info': {'end': '2018', 'full_name': 'Luis LÓPEZ GUERRA', 'start': '2008'},
              'role': 'judge'},
             {'name': 'SAJÓ', 'info': {'end': '2017', 'full_name': 'András SAJÓ', 'start': '2008'}, 'role': 'judge'},
             {'name': 'KARAKAŞ', 'info': {'end': '2019', 'full_name': 'Işıl KARAKAŞ', 'start': '2008'},
              'role': 'judge'},
             {'name': 'MØSE', 'info': {'end': '2018', 'full_name': 'Erik MØSE', 'start': '2011'}, 'role': 'judge'},
             {'name': 'PEJCHAL', 'info': {'end': None, 'full_name': 'Aleš PEJCHAL', 'start': '2012'}, 'role': 'judge'},
             {'name': 'WOJTYCZEK', 'info': {'end': None, 'full_name': 'Krzysztof WOJTYCZEK', 'start': '2012'},
              'role': 'judge'},
             {'name': 'KŪRIS', 'info': {'end': '2004', 'full_name': 'Pranas KŪRIS', 'start': '1994'}, 'role': 'judge'},
             {'name': 'MITS', 'info': {'end': None, 'full_name': 'Mārtiņš MITS', 'start': '2015'}, 'role': 'judge'},
             {'name': 'RAVARANI', 'info': {'end': None, 'full_name': 'Georges RAVARANI', 'start': '2015'},
              'role': 'judge'},
             {'name': 'PASTOR VILANOVA', 'info': {'end': None, 'full_name': 'Pere PASTOR VILANOVA', 'start': '2015'},
              'role': 'judge'},
             {'name': 'POLÁČKOVÁ', 'info': {'end': None, 'full_name': 'Alena POLÁČKOVÁ', 'start': '2016'},
              'role': 'judge'},
             {'name': 'SERGHIDES', 'info': {'end': None, 'full_name': 'Georgios SERGHIDES', 'start': '2016'},
              'role': 'judge'}]
        ]

        return {'docs': docs, 'doc_ids': doc_ids, 'build': build, 'output': outputs}

    @staticmethod
    def test_parse_document(prep):
        with patch('echr.steps.preprocess_documents.load_judges_info') as judge, open(
                'tests/data/judges_per_country.json', 'r') as f:
            judge.return_value = json.load(f)

        for i in range(len(prep['doc_ids'])):
            parsed, _, _ = parse_document(prep['docs'][i], prep['doc_ids'][i], prep['build'])
            res = parsed['decision_body']
            assert res == prep['output'][i]
